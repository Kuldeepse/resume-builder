import base64
import ipaddress
import io
import json
import math
import os
import re
from html import escape
from typing import Any, Optional
from urllib.parse import urlparse

from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from google import genai
from google.genai import types
from PyPDF2 import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from supabase import Client, create_client

app = FastAPI()

FRONTEND_ORIGINS = [
    origin.strip()
    for origin in os.getenv(
        "FRONTEND_ORIGINS",
        "https://cognitwistai.duckdns.org,http://localhost:3000",
    ).split(",")
    if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=FRONTEND_ORIGINS,
    allow_credentials=False,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["Content-Type", "Accept"],
)

supabase_url = os.getenv("SUPABASE_URL") or "https://supabase.co"
supabase_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY") or os.getenv("SUPABASE_KEY") or "placeholder-key"
supabase: Client = create_client(supabase_url, supabase_key)
gemini_client = genai.Client()

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_PROFILE_CHARS = 30000
ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_UPLOAD_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
}
ALLOWED_JOB_DOMAINS = {
    "linkedin.com",
    "indeed.com",
    "lever.co",
    "greenhouse.io",
    "workdayjobs.com",
    "myworkdayjobs.com",
    "smartrecruiters.com",
    "ashbyhq.com",
    "workable.com",
}


def safe_text(value: Any, limit: int = 4000) -> str:
    return str(value or "").strip()[:limit]


def safe_list(value: Any, limit: int = 50, item_limit: int = 1000) -> list[str]:
    if not isinstance(value, list):
        return []
    return [safe_text(item, item_limit) for item in value[:limit] if safe_text(item, item_limit)]


def safe_pdf_text(value: Any) -> str:
    return escape(str(value or ""), quote=True)


def safe_filename(value: str, suffix: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._-")[:80]
    return f"{normalized or 'candidate'}_{suffix}"


def safe_job_url(value: Any) -> Optional[str]:
    candidate = safe_text(value, 2000)
    if not candidate:
        return None

    try:
        parsed = urlparse(candidate)
    except ValueError:
        return None

    if parsed.scheme.lower() != "https" or not parsed.hostname or parsed.username or parsed.password:
        return None

    hostname = parsed.hostname.lower().rstrip(".")
    try:
        ipaddress.ip_address(hostname)
        return None
    except ValueError:
        pass

    if not any(hostname == domain or hostname.endswith(f".{domain}") for domain in ALLOWED_JOB_DOMAINS):
        return None
    return candidate


async def read_upload(upload: UploadFile) -> tuple[bytes, str]:
    filename = upload.filename or ""
    extension = os.path.splitext(filename.lower())[1]
    if extension not in ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are accepted.")
    if upload.content_type not in ALLOWED_UPLOAD_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported upload content type.")

    file_bytes = await upload.read(MAX_UPLOAD_BYTES + 1)
    if len(file_bytes) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Uploaded CV must be 5 MB or smaller.")
    if not file_bytes:
        raise HTTPException(status_code=400, detail="The uploaded CV is empty.")
    return file_bytes, extension


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Unable to read the PDF CV: {exc}") from exc


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    blocks.append(row_text)
        return "\n".join(blocks).strip()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Unable to read the DOCX CV: {exc}") from exc


async def extract_resume_text(upload: UploadFile) -> str:
    file_bytes, extension = await read_upload(upload)
    extracted = extract_pdf_text(file_bytes) if extension == ".pdf" else extract_docx_text(file_bytes)
    extracted = re.sub(r"\n{3,}", "\n\n", extracted).strip()
    if len(extracted) < 80:
        raise HTTPException(
            status_code=400,
            detail="Very little text could be extracted. Scanned/image-only PDFs are not supported; use a text PDF, DOCX, or paste the CV.",
        )
    return extracted[:MAX_PROFILE_CHARS]


def normalise_resume(value: Any, full_name: str, linkedin_profile: Optional[str]) -> dict[str, Any]:
    raw = value if isinstance(value, dict) else {}
    contact = raw.get("contact") if isinstance(raw.get("contact"), dict) else {}
    if linkedin_profile and not contact.get("linkedin"):
        contact["linkedin"] = linkedin_profile

    experience = []
    for item in raw.get("experience", []) if isinstance(raw.get("experience"), list) else []:
        if not isinstance(item, dict):
            continue
        experience.append(
            {
                "company": safe_text(item.get("company"), 300),
                "role": safe_text(item.get("role"), 300),
                "duration": safe_text(item.get("duration"), 200),
                "bullet_points": safe_list(item.get("bullet_points"), 12, 1200),
            }
        )

    education = []
    for item in raw.get("education", []) if isinstance(raw.get("education"), list) else []:
        if isinstance(item, dict):
            education.append(
                {
                    "qualification": safe_text(item.get("qualification"), 300),
                    "institution": safe_text(item.get("institution"), 300),
                    "duration": safe_text(item.get("duration"), 200),
                }
            )

    projects = []
    for item in raw.get("projects", []) if isinstance(raw.get("projects"), list) else []:
        if isinstance(item, dict):
            projects.append(
                {
                    "name": safe_text(item.get("name"), 300),
                    "description": safe_text(item.get("description"), 1500),
                    "impact": safe_text(item.get("impact"), 800),
                }
            )

    return {
        "full_name": safe_text(raw.get("full_name") or full_name, 200),
        "headline": safe_text(raw.get("headline"), 300),
        "contact": {
            "email": safe_text(contact.get("email"), 250),
            "phone": safe_text(contact.get("phone"), 100),
            "location": safe_text(contact.get("location"), 250),
            "linkedin": safe_text(contact.get("linkedin"), 500),
        },
        "professional_summary": safe_text(raw.get("professional_summary"), 2500),
        "skills": safe_list(raw.get("skills"), 60, 200),
        "experience": experience[:20],
        "education": education[:12],
        "certifications": safe_list(raw.get("certifications"), 30, 400),
        "projects": projects[:12],
        "achievements": safe_list(raw.get("achievements"), 30, 1000),
    }


def build_pdf_bytes(resume: dict[str, Any], target_role: str) -> bytes:
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=32, bottomMargin=32)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TStyle", parent=styles["Heading1"], fontSize=21, leading=25, spaceAfter=5)
    headline_style = ParagraphStyle("HStyle", parent=styles["Normal"], fontSize=10, leading=14, textColor=colors.HexColor("#334155"))
    section_style = ParagraphStyle(
        "SStyle",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceBefore=9,
        spaceAfter=4,
        textColor=colors.HexColor("#0F766E"),
    )
    body_style = ParagraphStyle("BStyle", parent=styles["Normal"], fontSize=9.5, leading=13)

    story = [Paragraph(f"<b>{safe_pdf_text(resume.get('full_name'))}</b>", title_style)]
    headline = resume.get("headline") or target_role
    story.append(Paragraph(safe_pdf_text(headline), headline_style))

    contact = resume.get("contact", {})
    contact_line = " | ".join(
        safe_text(contact.get(field), 500)
        for field in ("email", "phone", "location", "linkedin")
        if safe_text(contact.get(field), 500)
    )
    if contact_line:
        story.append(Paragraph(safe_pdf_text(contact_line), headline_style))
    story.append(Spacer(1, 5))

    def add_section(title: str, values: list[str]) -> None:
        clean_values = [item for item in values if item]
        if not clean_values:
            return
        story.append(Paragraph(f"<b>{safe_pdf_text(title)}</b>", section_style))
        for item in clean_values:
            story.append(Paragraph(safe_pdf_text(item), body_style))

    add_section("Professional Summary", [safe_text(resume.get("professional_summary"), 5000)])
    skills = resume.get("skills", [])
    add_section("Core Skills", [", ".join(skills)] if isinstance(skills, list) else [])

    experience = resume.get("experience", [])
    if experience:
        story.append(Paragraph("<b>Professional Experience</b>", section_style))
        for item in experience:
            heading = f"<b>{safe_pdf_text(item.get('role'))}</b> — {safe_pdf_text(item.get('company'))}"
            duration = safe_text(item.get("duration"), 200)
            if duration:
                heading += f" ({safe_pdf_text(duration)})"
            story.append(Paragraph(heading, body_style))
            for bullet in item.get("bullet_points", []):
                story.append(Paragraph(f"• {safe_pdf_text(bullet)}", body_style))
            story.append(Spacer(1, 3))

    education_lines = [
        " — ".join(filter(None, [safe_text(item.get("qualification")), safe_text(item.get("institution")), safe_text(item.get("duration"))]))
        for item in resume.get("education", [])
    ]
    add_section("Education", education_lines)
    add_section("Certifications", [f"• {item}" for item in resume.get("certifications", [])])

    project_lines = []
    for item in resume.get("projects", []):
        line = f"{safe_text(item.get('name'))}: {safe_text(item.get('description'), 2000)}"
        if safe_text(item.get("impact"), 1000):
            line += f" Impact: {safe_text(item.get('impact'), 1000)}"
        project_lines.append(line)
    add_section("Selected Projects", project_lines)
    add_section("Achievements", [f"• {item}" for item in resume.get("achievements", [])])

    doc.build(story)
    return output.getvalue()


def build_docx_bytes(resume: dict[str, Any], target_role: str) -> bytes:
    document = Document()
    document.add_heading(safe_text(resume.get("full_name"), 200), level=0)
    document.add_paragraph(safe_text(resume.get("headline") or target_role, 300))

    contact = resume.get("contact", {})
    contact_line = " | ".join(
        safe_text(contact.get(field), 500)
        for field in ("email", "phone", "location", "linkedin")
        if safe_text(contact.get(field), 500)
    )
    if contact_line:
        document.add_paragraph(contact_line)

    def add_heading_and_paragraph(title: str, text: str) -> None:
        if text:
            document.add_heading(title, level=1)
            document.add_paragraph(text)

    add_heading_and_paragraph("Professional Summary", safe_text(resume.get("professional_summary"), 5000))
    skills = resume.get("skills", [])
    add_heading_and_paragraph("Core Skills", ", ".join(skills) if isinstance(skills, list) else "")

    if resume.get("experience"):
        document.add_heading("Professional Experience", level=1)
        for item in resume["experience"]:
            document.add_heading(
                " — ".join(filter(None, [safe_text(item.get("role")), safe_text(item.get("company")), safe_text(item.get("duration"))])),
                level=2,
            )
            for bullet in item.get("bullet_points", []):
                document.add_paragraph(bullet, style="List Bullet")

    if resume.get("education"):
        document.add_heading("Education", level=1)
        for item in resume["education"]:
            document.add_paragraph(
                " — ".join(filter(None, [safe_text(item.get("qualification")), safe_text(item.get("institution")), safe_text(item.get("duration"))]))
            )

    if resume.get("certifications"):
        document.add_heading("Certifications", level=1)
        for item in resume["certifications"]:
            document.add_paragraph(item, style="List Bullet")

    if resume.get("projects"):
        document.add_heading("Selected Projects", level=1)
        for item in resume["projects"]:
            document.add_heading(safe_text(item.get("name")), level=2)
            document.add_paragraph(safe_text(item.get("description"), 2000))
            if safe_text(item.get("impact"), 1000):
                document.add_paragraph(f"Impact: {safe_text(item.get('impact'), 1000)}")

    if resume.get("achievements"):
        document.add_heading("Achievements", level=1)
        for item in resume["achievements"]:
            document.add_paragraph(item, style="List Bullet")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def interview_distribution(interview_type: str, requested_count: int) -> str:
    if interview_type == "hr":
        return f"Generate exactly {requested_count} HR screening questions covering motivation, fit, communication, availability and expectations."
    if interview_type == "behavioural":
        return f"Generate exactly {requested_count} behavioural questions and evidence-based STAR responses."
    return f"Generate exactly {requested_count} technical questions covering architecture, controls, trade-offs, delivery risk and operational readiness."


@app.get("/health/")
@app.get("/health")
async def health_check():
    return {"status": "healthy"}


@app.post("/build-resume/")
@app.post("/build-resume")
async def build_and_compare_resume(
    full_name: str = Form(...),
    target_role: str = Form(...),
    job_description: str = Form(...),
    career_history: Optional[str] = Form(""),
    resume_file: Optional[UploadFile] = File(None),
    linkedin_profile: Optional[str] = Form(None),
    interview_duration: Any = Form("30 minutes"),
    total_questions_requested: Any = Form(5),
    interview_type: Optional[str] = Form("technical"),
):
    pasted_profile = safe_text(career_history, MAX_PROFILE_CHARS)
    uploaded_profile = await extract_resume_text(resume_file) if resume_file else ""
    source_profile = "\n\n".join(part for part in (uploaded_profile, pasted_profile) if part).strip()
    if not source_profile:
        raise HTTPException(status_code=400, detail="Upload a CV or paste the candidate’s career history.")

    try:
        requested_count = max(5, min(25, int(total_questions_requested)))
    except (ValueError, TypeError):
        requested_count = 5

    current_type = safe_text(interview_type or "technical", 30).lower()
    if current_type not in {"hr", "behavioural", "technical"}:
        current_type = "technical"

    system_prompt = f"""You are an expert UK recruiter, CV writer and ATS analyst.
Return one valid JSON object only.

NON-NEGOTIABLE EVIDENCE RULES:
- Use only employers, roles, dates, qualifications, certifications, technologies, responsibilities and metrics explicitly present in the candidate source.
- Never invent or infer an achievement as fact.
- You may reorder, condense and rewrite verified evidence to improve relevance and clarity.
- If the job description asks for unsupported experience, list it under missing_skills or evidence_warnings; do not add it to the CV.
- Preserve the candidate's seniority and distinguish personal delivery from team activity.

REQUIRED JSON SCHEMA:
{{
  "match_score": 75,
  "matched_requirements": ["verified requirement"],
  "missing_skills": ["unsupported or missing requirement"],
  "evidence_warnings": ["claim that needs candidate confirmation"],
  "tailoring_tips": ["specific action"],
  "change_log": [
    {{
      "section": "Professional Summary",
      "original": "short source extract",
      "revised": "proposed wording",
      "reason": "why it improves fit",
      "evidence_status": "verified|needs_confirmation"
    }}
  ],
  "tell_me_about_yourself": "role-aligned evidence-based introduction",
  "interview_questions": [{{"question": "string", "response": "structured response using only verified evidence"}}],
  "follow_up_questions": ["candidate question for interviewer"],
  "resume": {{
    "full_name": "string",
    "headline": "string",
    "contact": {{"email": "", "phone": "", "location": "", "linkedin": ""}},
    "professional_summary": "string",
    "skills": ["skill"],
    "experience": [{{"company": "", "role": "", "duration": "", "bullet_points": ["verified achievement"]}}],
    "education": [{{"qualification": "", "institution": "", "duration": ""}}],
    "certifications": ["verified certification"],
    "projects": [{{"name": "", "description": "", "impact": ""}}],
    "achievements": ["verified achievement"]
  }}
}}

INTERVIEW REQUIREMENT:
{interview_distribution(current_type, requested_count)}
For behavioural responses use Situation, Task, Action and Result. For technical responses explain architecture, choices, controls, validation and outcome. For HR responses keep the answer concise and credible.
Generate 3 to 5 intelligent questions for the candidate to ask the interviewer."""

    linkedin_context = f"\nCandidate LinkedIn URL: {safe_text(linkedin_profile, 500)}" if linkedin_profile else ""

    try:
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=(
                f"Candidate Name: {safe_text(full_name, 200)}\n"
                f"Target Role: {safe_text(target_role, 300)}{linkedin_context}\n"
                f"Interview Duration: {safe_text(interview_duration, 50)}\n"
                f"Candidate Source:\n{source_profile}\n\n"
                f"Job Description:\n{safe_text(job_description, MAX_PROFILE_CHARS)}"
            ),
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                response_mime_type="application/json",
                temperature=0.1,
            ),
        )
        raw_text = (response.text or "").strip()
        if "```json" in raw_text:
            raw_text = raw_text.split("```json")[-1].split("```")[0].strip()
        elif "```" in raw_text:
            raw_text = raw_text.split("```")[-1].split("```")[0].strip()
        analysis_result = json.loads(raw_text)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"CV tailoring failed: {exc}") from exc

    resume = normalise_resume(analysis_result.get("resume"), full_name, linkedin_profile)
    pdf_bytes = build_pdf_bytes(resume, target_role)
    docx_bytes = build_docx_bytes(resume, target_role)
    base_name = safe_filename(f"{full_name}_{target_role}", "Tailored_CV")

    raw_questions = analysis_result.get("interview_questions", [])
    final_questions = []
    if isinstance(raw_questions, list):
        for item in raw_questions[:requested_count]:
            if isinstance(item, dict):
                final_questions.append(
                    {
                        "question": safe_text(item.get("question"), 1200),
                        "response": safe_text(item.get("response"), 6000),
                    }
                )

    change_log = []
    for item in analysis_result.get("change_log", []) if isinstance(analysis_result.get("change_log"), list) else []:
        if isinstance(item, dict):
            change_log.append(
                {
                    "section": safe_text(item.get("section"), 200),
                    "original": safe_text(item.get("original"), 2000),
                    "revised": safe_text(item.get("revised"), 3000),
                    "reason": safe_text(item.get("reason"), 1200),
                    "evidence_status": "needs_confirmation"
                    if safe_text(item.get("evidence_status"), 50) == "needs_confirmation"
                    else "verified",
                }
            )

    return {
        "match_score": max(0, min(100, int(analysis_result.get("match_score", 70)))),
        "matched_requirements": safe_list(analysis_result.get("matched_requirements"), 40, 500),
        "missing_skills": safe_list(analysis_result.get("missing_skills"), 40, 500),
        "evidence_warnings": safe_list(analysis_result.get("evidence_warnings"), 30, 1000),
        "tailoring_tips": safe_list(analysis_result.get("tailoring_tips"), 30, 1000),
        "change_log": change_log[:30],
        "tell_me_about_yourself": safe_text(analysis_result.get("tell_me_about_yourself"), 7000),
        "interview_questions": final_questions,
        "follow_up_questions": safe_list(analysis_result.get("follow_up_questions"), 10, 1200),
        "resume": resume,
        "source_preview": source_profile[:2500],
        "pdf_base64": base64.b64encode(pdf_bytes).decode("ascii"),
        "pdf_filename": f"{base_name}.pdf",
        "docx_base64": base64.b64encode(docx_bytes).decode("ascii"),
        "docx_filename": f"{base_name}.docx",
        "shareable_url": None,
    }


@app.post("/search-jobs/")
@app.post("/search-jobs")
async def search_jobs(
    target_role: str = Form(...),
    location_city: str = Form(...),
    resume_skills: Optional[str] = Form(None),
    resume_file: Optional[UploadFile] = File(None),
):
    uploaded_profile = await extract_resume_text(resume_file) if resume_file else ""
    candidate_profile = uploaded_profile or safe_text(resume_skills, MAX_PROFILE_CHARS)
    if not candidate_profile:
        raise HTTPException(status_code=400, detail="Upload a CV or enter a role and skills summary.")

    search_query = (
        f'"{safe_text(target_role, 300)}" jobs in "{safe_text(location_city, 200)}" posted last 14 days '
        "site:linkedin.com OR site:indeed.com OR site:lever.co OR site:greenhouse.io "
        "OR site:myworkdayjobs.com OR site:smartrecruiters.com OR site:jobs.ashbyhq.com"
    )

    try:
        search_response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=(
                "Find current, real job openings. Return useful result text with title, company, location, "
                f"description, date and exact application URL. Query: {search_query}"
            ),
            config=types.GenerateContentConfig(
                tools=[types.Tool(google_search=types.GoogleSearch())],
                temperature=0.1,
            ),
        )
        raw_web_data = getattr(search_response, "text", "") or str(search_response)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Live job search failed: {exc}") from exc

    system_prompt = """You are a job-matching extraction engine.
Return valid JSON only. Use only jobs present in the supplied grounded search material. Never fabricate a job or link.

SCHEMA:
{
  "jobs": [
    {
      "title": "",
      "company": "",
      "location": "",
      "salary": "Not disclosed",
      "posted": "",
      "description": "concise evidence-based role description",
      "skills": ["skill"],
      "link": "exact HTTPS application URL",
      "match_score": 0,
      "matched_requirements": ["match"],
      "missing_requirements": ["gap"],
      "recommendation": "Apply|Apply after tailoring|Review carefully"
    }
  ],
  "best_match_summary": ""
}

Score each job against the candidate profile, role objective and location. Do not treat unsupported experience as a match."""

    try:
        formatting_response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=(
                f"Grounded Job Search Material:\n{raw_web_data}\n\n"
                f"Candidate Profile:\n{candidate_profile[:MAX_PROFILE_CHARS]}\n\n"
                f"Role Objective: {safe_text(target_role, 300)}\n"
                f"Location: {safe_text(location_city, 200)}"
            ),
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                response_mime_type="application/json",
                temperature=0.1,
            ),
        )
        clean_text = (formatting_response.text or "").strip()
        if "```json" in clean_text:
            clean_text = clean_text.split("```json")[-1].split("```")[0].strip()
        elif "```" in clean_text:
            clean_text = clean_text.split("```")[-1].split("```")[0].strip()
        jobs_result = json.loads(clean_text)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Job result extraction failed: {exc}") from exc

    sanitized_jobs = []
    raw_jobs = jobs_result.get("jobs", []) if isinstance(jobs_result, dict) else []
    for item in raw_jobs if isinstance(raw_jobs, list) else []:
        if not isinstance(item, dict):
            continue
        validated_link = safe_job_url(item.get("link"))
        try:
            match_score = max(0, min(100, int(item.get("match_score", 0))))
        except (ValueError, TypeError):
            match_score = 0
        recommendation = safe_text(item.get("recommendation"), 120)
        if recommendation not in {"Apply", "Apply after tailoring", "Review carefully"}:
            recommendation = "Review carefully"
        sanitized_jobs.append(
            {
                "title": safe_text(item.get("title") or "Role", 240),
                "company": safe_text(item.get("company") or "Company", 240),
                "location": safe_text(item.get("location") or location_city, 240),
                "salary": safe_text(item.get("salary") or "Not disclosed", 240),
                "posted": safe_text(item.get("posted"), 120),
                "description": safe_text(item.get("description"), 5000),
                "skills": safe_list(item.get("skills"), 30, 160),
                "link": validated_link or "search on company website",
                "match_score": match_score,
                "matched_requirements": safe_list(item.get("matched_requirements"), 20, 500),
                "missing_requirements": safe_list(item.get("missing_requirements"), 20, 500),
                "recommendation": recommendation,
            }
        )

    sanitized_jobs.sort(key=lambda job: job["match_score"], reverse=True)
    return {
        "jobs": sanitized_jobs[:40],
        "best_match_summary": safe_text(
            jobs_result.get("best_match_summary") if isinstance(jobs_result, dict) else "",
            2000,
        ),
        "profile_preview": candidate_profile[:1800],
    }
